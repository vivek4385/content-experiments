from anthropic import Anthropic
import requests
import xml.etree.ElementTree as ET
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re

def fetch_sitemap(sitemap_url):
    """
    Fetch and parse XML sitemap to extract URLs and titles.
    
    Returns:
        list: List of dicts with 'url' and 'title' keys
    """
    try:
        response = requests.get(sitemap_url, timeout=10)
        response.raise_for_status()
        
        root = ET.fromstring(response.content)
        
        # Handle XML namespaces
        namespaces = {'ns': 'http://www.sitemaps.org/schemas/sitemap/0.9'}
        
        urls = []
        for url in root.findall('ns:url', namespaces):
            loc = url.find('ns:loc', namespaces)
            if loc is not None:
                url_text = loc.text
                # Extract title from URL (last segment)
                title = url_text.rstrip('/').split('/')[-1].replace('-', ' ').title()
                urls.append({
                    'url': url_text,
                    'title': title
                })
        
        return urls
    
    except Exception as e:
        raise Exception(f"Failed to fetch sitemap: {str(e)}")


def add_internal_links(article_text, sitemap_url, num_links, priority_urls, api_key, progress_callback=None):
    """
    Add internal links to an article using Claude AI.
    
    Args:
        article_text: The article content to add links to
        sitemap_url: URL to the XML sitemap
        num_links: Number of internal links to add
        priority_urls: Comma-separated list of priority URLs (or empty string)
        api_key: Anthropic API key
        progress_callback: Optional function to call with progress updates
    
    Returns:
        Document: Word document with hyperlinks added
    """
    
    def update_progress(text):
        if progress_callback:
            progress_callback(text)
        else:
            print(text)
    
    # Initialize client
    client = Anthropic(api_key=api_key)
    
    # Fetch sitemap
    update_progress("Fetching sitemap...")
    sitemap_pages = fetch_sitemap(sitemap_url)
    
    # Format sitemap for prompt
    sitemap_text = "\n".join([f"- {page['title']}: {page['url']}" for page in sitemap_pages])
    
    # Parse priority URLs
    priority_urls_list = []
    if priority_urls.strip():
        priority_urls_list = [url.strip() for url in priority_urls.replace('\n', ',').split(',') if url.strip()]
    
    priority_text = "\n".join([f"- {url}" for url in priority_urls_list]) if priority_urls_list else "None specified"
    
    # Create prompt
    update_progress("Analyzing article for link opportunities...")
    
    prompt = f"""You are an internal linking specialist. Add internal links to this article.

ARTICLE TO ADD LINKS TO:
{article_text}

AVAILABLE PAGES (from sitemap):
{sitemap_text}

PRIORITY URLS (use these first if contextually relevant):
{priority_text}

NUMBER OF LINKS TO ADD: {num_links}

INSTRUCTIONS:
1. Read the article carefully and identify {num_links} opportunities for internal links
2. PRIORITIZE the priority URLs first - if they fit contextually, use them before other pages
3. For each link:
   - Find natural anchor text (2-5 words) that describes the destination page
   - Choose the most relevant page from the available pages
   - Ensure the link fits naturally in the sentence context
   - Do NOT force links where they don't make sense
4. Each URL should be used ONLY ONCE across the entire article
5. Distribute links throughout the article - avoid clustering in one section
6. Links should feel natural, not keyword-stuffed

OUTPUT FORMAT:
Return the article with links in this EXACT format:
- Use double square brackets for links: [[anchor text|URL]]
- Example: "This is a sentence about [[payment processing|https://example.com/payments]] in the article."
- Do NOT use markdown format like [text](url)
- Do NOT use HTML format like <a href="">
- Use ONLY the [[anchor text|URL]] format

CRITICAL RULES:
- Use each URL only once
- Add exactly {num_links} links (or fewer if not enough good opportunities)
- Links must be contextually relevant
- Maintain all original article content and structure
- Only add the link syntax, don't modify any other text

Return the complete article with internal links added:"""

    # Call Claude
    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=8000,
        messages=[{
            "role": "user",
            "content": prompt
        }]
    )
    
    linked_article_text = message.content[0].text.strip()
    
    # Create Word document with hyperlinks
    update_progress("Creating Word document with hyperlinks...")
    
    doc = Document()
    
    # Parse the linked article and add to document
    # Split by paragraphs
    paragraphs = linked_article_text.split('\n\n')
    
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        
        # Add paragraph
        paragraph = doc.add_paragraph()
        
        # Find all links in format [[anchor text|URL]]
        parts = re.split(r'\[\[([^\]]+)\]\]', para_text)
        
        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Regular text
                if part:
                    paragraph.add_run(part)
            else:
                # Link part - format: anchor text|URL
                if '|' in part:
                    anchor, url = part.split('|', 1)
                    anchor = anchor.strip()
                    url = url.strip()
                    
                    # Add hyperlink
                    add_hyperlink(paragraph, url, anchor)
                else:
                    # Malformed link, just add as text
                    paragraph.add_run(f"[[{part}]]")
    
    update_progress("✓ Document created with hyperlinks!")
    
    return doc


def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph in Word document.
    
    Args:
        paragraph: The paragraph object
        url: The URL to link to
        text: The anchor text
    """
    # Create hyperlink element
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create run element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # Add blue color and underline
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)